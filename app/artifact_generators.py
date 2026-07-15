"""
Document generators for Module 5 (FR-5.3). Each generator returns
(file_bytes, filename, mime_type) given a system and the live database
session. Lazy imports keep the app importable even before python-docx
and openpyxl are installed.
"""
import io
import json
from datetime import datetime, timezone


def _now_str():
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


# ---------------------------------------------------------------------------
# .docx generators
# ---------------------------------------------------------------------------
def _new_doc():
    from docx import Document
    return Document()


def _doc_disclaimer(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.italic = True
        run.font.size = None  # default


def gen_ssp(system, controls, poams, inv, **_):
    """System Security Plan (.docx) - FR-5.1, FR-2.3.6, FR-5.5."""
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    title = doc.add_heading(f"System Security Plan", level=0)
    doc.add_heading(f"{system.name} ({system.acronym or '-'})", level=1)
    p = doc.add_paragraph()
    p.add_run("UNCLASSIFIED. LOCAL VALIDATION REQUIRED. ").bold = True
    p.add_run(f"Generated {_now_str()} by RMF ATO Navigator prototype.")

    # 1. Identification
    doc.add_heading("1. System Identification", level=1)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"
    for k, v in [
        ("System name", system.name),
        ("Acronym", system.acronym or "-"),
        ("System type", system.system_type or "-"),
        ("ATO type", system.ato_type or "-"),
        ("Current RMF step", str(system.current_step)),
        ("Impact level (FIPS 199)", system.impact_level or "-"),
        ("CIA (C/I/A)",
         f"{system.cia_confidentiality or '?'}/{system.cia_integrity or '?'}/{system.cia_availability or '?'}"),
        ("National Security System", "Yes" if system.is_nss else "No"),
        ("Data types", system.data_types or "-"),
    ]:
        row = tbl.add_row().cells
        row[0].text = k
        row[1].text = str(v)

    # 2. Mission / Boundary
    doc.add_heading("2. Mission and Authorization Boundary", level=1)
    doc.add_heading("Mission description", level=2)
    doc.add_paragraph(system.mission_desc or "[To be documented]")
    doc.add_heading("Authorization boundary", level=2)
    doc.add_paragraph(system.boundary_statement or "[To be documented]")

    # 3. Inventory summary
    doc.add_heading("3. Component Inventory Summary", level=1)
    counts = {"hardware": 0, "software": 0, "firmware": 0, "sbom": 0}
    for it in inv:
        counts[it.item_type] = counts.get(it.item_type, 0) + 1
    inv_tbl = doc.add_table(rows=1, cols=2)
    inv_tbl.style = "Light Grid Accent 1"
    inv_tbl.rows[0].cells[0].text = "Inventory class"
    inv_tbl.rows[0].cells[1].text = "Count"
    for k, v in counts.items():
        r = inv_tbl.add_row().cells
        r[0].text = k.capitalize()
        r[1].text = str(v)
    doc.add_paragraph(
        "Detailed inventory exists in the Inventory module and is referenced "
        "by control CM-8.")

    # 4. Control Implementation Statements
    doc.add_heading("4. Control Implementation Statements", level=1)
    if not controls:
        doc.add_paragraph(
            "[No controls have been applied to this system yet. Apply baseline "
            "controls in the per-system control workspace.]")
    for c in controls:
        doc.add_heading(f"{c.nist_id} - {c.title or ''}", level=2)
        meta = doc.add_paragraph()
        meta.add_run(f"Status: {c.implementation_status or 'not_implemented'}  |  ").bold = True
        meta.add_run(f"Type: {c.control_type or '-'}  |  ")
        meta.add_run(f"Family: {c.family or '-'}  |  ")
        meta.add_run(f"Baseline: {c.baseline or '-'}")
        doc.add_paragraph(c.implementation_statement
                          or "[Implementation statement to be drafted.]")

    # 5. POA&M Reference
    doc.add_heading("5. Plan of Action and Milestones (Reference)", level=1)
    if not poams:
        doc.add_paragraph("No POA&M entries on register.")
    else:
        ptbl = doc.add_table(rows=1, cols=5)
        ptbl.style = "Light Grid Accent 1"
        head = ptbl.rows[0].cells
        head[0].text = "POA&M ID"
        head[1].text = "Control"
        head[2].text = "Severity"
        head[3].text = "Status"
        head[4].text = "Weakness"
        for p in poams:
            r = ptbl.add_row().cells
            r[0].text = p.poam_id[:8]
            r[1].text = p.control_nist_id or "-"
            r[2].text = p.severity or "-"
            r[3].text = p.status or "-"
            r[4].text = (p.weakness or "")[:120]

    # 6. Approval
    doc.add_heading("6. Approval", level=1)
    doc.add_paragraph(
        "ISSO: ____________________________________________  Date: ___________")
    doc.add_paragraph(
        "ISSM: ____________________________________________  Date: ___________")
    doc.add_paragraph(
        "AO:   ____________________________________________  Date: ___________")

    buf = io.BytesIO()
    doc.save(buf)
    return (buf.getvalue(),
            f"SSP_{system.acronym or 'system'}_{datetime.now().strftime('%Y%m%d')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def gen_stakeholder_matrix(system, users, **_):
    from docx import Document
    doc = Document()
    doc.add_heading(f"Stakeholder Matrix - {system.name}", level=0)
    doc.add_paragraph(f"Generated {_now_str()} by RMF ATO Navigator prototype.")
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    head = tbl.rows[0].cells
    head[0].text = "Role"
    head[1].text = "Name"
    head[2].text = "Email"
    head[3].text = "Approval authority"
    for role, uid_attr in [
        ("ISSO", "isso_id"), ("ISSM", "issm_id"),
        ("SCA", "sca_id"), ("AO / Rep", "ao_id"),
    ]:
        u = users.get(getattr(system, uid_attr))
        r = tbl.add_row().cells
        r[0].text = role
        r[1].text = u.name if u else "[unassigned]"
        r[2].text = u.email if u else "-"
        r[3].text = ("Authorization decision" if role.startswith("AO") else
                     "Closure approval" if role == "ISSM" else
                     "Independent assessment" if role == "SCA" else
                     "Daily RMF execution")
    doc.add_paragraph(
        "ISSO drafts; ISSM approves before Step 1.")
    buf = io.BytesIO()
    doc.save(buf)
    return (buf.getvalue(),
            f"StakeholderMatrix_{system.acronym or 'system'}_{datetime.now().strftime('%Y%m%d')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def gen_fips199(system, **_):
    from docx import Document
    doc = Document()
    doc.add_heading(f"FIPS 199 Categorization Worksheet - {system.name}", level=0)
    doc.add_paragraph(f"Generated {_now_str()} by RMF ATO Navigator prototype.")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light Grid Accent 1"
    head = tbl.rows[0].cells
    head[0].text = "Dimension"
    head[1].text = "Rating"
    for k, v in [
        ("Confidentiality", system.cia_confidentiality or "-"),
        ("Integrity", system.cia_integrity or "-"),
        ("Availability", system.cia_availability or "-"),
        ("Overall impact", system.impact_level or "-"),
    ]:
        r = tbl.add_row().cells
        r[0].text = k
        r[1].text = v
    doc.add_heading("Data types processed", level=1)
    doc.add_paragraph(system.data_types or "[None recorded]")
    doc.add_heading("Categorization rationale", level=1)
    doc.add_paragraph(
        f"Mission impact analysis indicates that loss of {system.cia_confidentiality or '?'} "
        f"confidentiality, {system.cia_integrity or '?'} integrity, and "
        f"{system.cia_availability or '?'} availability would have the corresponding effect "
        f"on the supported mission. Overall impact level set at {system.impact_level or '-'}.")
    doc.add_paragraph(
        "Refine the rationale narrative before submission. NIST SP 800-60 information "
        "type mapping must be attached.")
    buf = io.BytesIO()
    doc.save(buf)
    return (buf.getvalue(),
            f"FIPS199_{system.acronym or 'system'}_{datetime.now().strftime('%Y%m%d')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ---------------------------------------------------------------------------
# .xlsx generators
# ---------------------------------------------------------------------------
def _xlsx_header(ws, headers):
    from openpyxl.styles import Font, PatternFill, Alignment
    bold_white = Font(bold=True, color="FFFFFFFF")
    fill = PatternFill(start_color="FF1B3A6B", end_color="FF1B3A6B",
                        fill_type="solid")
    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = bold_white
        cell.fill = fill
        cell.alignment = Alignment(horizontal="left", vertical="center",
                                    wrap_text=True)
    ws.freeze_panes = "A2"


def _xlsx_autosize(ws, max_width=60):
    for col in ws.columns:
        m = 10
        col_letter = col[0].column_letter
        for cell in col:
            v = "" if cell.value is None else str(cell.value)
            for line in v.split("\n"):
                m = max(m, min(len(line) + 2, max_width))
        ws.column_dimensions[col_letter].width = m


def gen_inventory_xlsx(system, items, kind, **_):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = f"{kind.capitalize()} Inventory"
    headers = ["Name", "Vendor", "Version", "IP", "MAC", "OS", "OS Version",
               "Patch Level", "License", "Origin", "EOS Date", "Owner",
               "Location", "Status"]
    _xlsx_header(ws, headers)
    for it in items:
        ws.append([
            it.name, it.vendor or "", it.version or "",
            it.ip_address or "", it.mac_address or "",
            it.os or "", it.os_version or "", it.patch_level or "",
            it.license or "", it.component_origin or "",
            it.eos_date.isoformat() if it.eos_date else "",
            it.owner or "", it.location or "", it.status or ""])
    _xlsx_autosize(ws)
    buf = io.BytesIO()
    wb.save(buf)
    return (buf.getvalue(),
            f"{kind.capitalize()}Inventory_{system.acronym or 'system'}_{datetime.now().strftime('%Y%m%d')}.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


def gen_hardware_inventory(system, hw, **_):
    return gen_inventory_xlsx(system, hw, "hardware")


def gen_software_inventory(system, sw, **_):
    return gen_inventory_xlsx(system, sw, "software")


def gen_vuln_trend(system, scans, **_):
    from openpyxl import Workbook
    from openpyxl.chart import BarChart, Reference
    wb = Workbook()
    ws = wb.active
    ws.title = "Scan Trend"
    headers = ["Scan date", "Source", "Filename", "Critical", "High",
               "Medium", "Low", "Total"]
    _xlsx_header(ws, headers)
    if not scans:
        ws.append(["[No scans imported yet]", "", "", "", "", "", "", ""])
    else:
        for sc in scans:
            ws.append([
                sc.scan_date.strftime("%Y-%m-%d") if sc.scan_date else "",
                sc.source, sc.filename or "",
                sc.critical or 0, sc.high or 0,
                sc.medium or 0, sc.low or 0, sc.total or 0,
            ])
    _xlsx_autosize(ws)

    # Bar chart of severity over time
    if scans and len(scans) >= 1:
        chart = BarChart()
        chart.type = "col"
        chart.style = 11
        chart.title = "Vulnerability Findings by Scan"
        chart.y_axis.title = "Findings"
        chart.x_axis.title = "Scan"
        data_ref = Reference(ws, min_col=4, min_row=1,
                              max_col=7, max_row=1 + len(scans))
        cats_ref = Reference(ws, min_col=1, min_row=2, max_row=1 + len(scans))
        chart.add_data(data_ref, titles_from_data=True)
        chart.set_categories(cats_ref)
        ws.add_chart(chart, "J2")

    buf = io.BytesIO()
    wb.save(buf)
    return (buf.getvalue(),
            f"VulnScanTrend_{system.acronym or 'system'}_{datetime.now().strftime('%Y%m%d')}.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


def gen_poam_register(system, poams, **_):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "POA&M Register"
    headers = [
        "System ID", "System Name", "POAM ID", "Weakness",
        "Source of Discovery", "Control / CCI", "Severity (CAT)",
        "Risk Level", "Status", "Date Identified", "Scheduled Completion",
        "Resources Required", "Milestones",
        "Validation Method", "Residual Risk", "Comments", "Date Closed",
    ]
    _xlsx_header(ws, headers)
    for p in poams:
        try:
            ms = json.loads(p.milestones_json or "[]")
        except (ValueError, TypeError):
            ms = []
        ms_text = "; ".join(
            f"{m.get('title','')} (due {m.get('due','-')}"
            + (", done" if m.get("completed_at") else "")
            + ")"
            for m in ms
        )
        ws.append([
            system.acronym or system.system_id, system.name,
            p.poam_id, (p.weakness or "")[:400],
            p.source or "", p.control_nist_id or "",
            p.severity or "", p.risk_level or "", p.status or "",
            p.date_identified.isoformat() if p.date_identified else "",
            p.scheduled_completion.isoformat() if p.scheduled_completion else "",
            p.resources_required or "", ms_text,
            p.validation_method or "", p.residual_risk or "",
            (p.comments or "").replace("\n", " "),
            p.closure_approved_at.isoformat() if p.closure_approved_at else "",
        ])
    _xlsx_autosize(ws)
    buf = io.BytesIO()
    wb.save(buf)
    return (buf.getvalue(),
            f"POAMRegister_{system.acronym or 'system'}_{datetime.now().strftime('%Y%m%d')}.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


GENERATORS = {
    "ssp": gen_ssp,
    "stakeholder_matrix": gen_stakeholder_matrix,
    "fips199": gen_fips199,
    "hardware_inventory": gen_hardware_inventory,
    "software_inventory": gen_software_inventory,
    "vuln_trend": gen_vuln_trend,
    "poam_register": gen_poam_register,
}
